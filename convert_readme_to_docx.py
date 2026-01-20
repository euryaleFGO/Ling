#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
将 README.md 转换为 Word 文档 (docx)
用于 PPT 汇报和文档分享
"""
import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
except ImportError:
    print("❌ 错误: 缺少 python-docx 库")
    print("请安装: pip install python-docx")
    sys.exit(1)

def setup_document():
    """创建并设置文档格式"""
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)
    doc.styles['Normal'].paragraph_format.line_spacing = 1.5
    
    # 设置标题样式
    for i in range(1, 7):
        style = doc.styles[f'Heading {i}']
        style.font.name = '黑体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        style.font.size = Pt(20 - i * 2)
        style.font.bold = True
    
    return doc

def parse_markdown_line(line):
    """解析 Markdown 行，返回类型和内容"""
    line = line.rstrip()
    
    # 标题
    if line.startswith('# '):
        return 'h1', line[2:].strip()
    elif line.startswith('## '):
        return 'h2', line[3:].strip()
    elif line.startswith('### '):
        return 'h3', line[4:].strip()
    elif line.startswith('#### '):
        return 'h4', line[4:].strip()
    elif line.startswith('##### '):
        return 'h5', line[5:].strip()
    elif line.startswith('###### '):
        return 'h6', line[6:].strip()
    
    # 列表
    elif re.match(r'^[-*+] ', line):
        return 'ul', line[2:].strip()
    elif re.match(r'^\d+\. ', line):
        return 'ol', re.sub(r'^\d+\. ', '', line).strip()
    
    # 代码块
    elif line.startswith('```'):
        return 'code_block', None
    
    # 分隔线
    elif line.strip() == '---':
        return 'hr', None
    
    # 普通段落
    elif line.strip():
        return 'p', line.strip()
    
    # 空行
    else:
        return 'blank', None

def add_formatted_text(paragraph, text):
    """向段落添加格式化文本"""
    from docx.oxml.ns import nsdecls
    
    # 移除 emoji（可选）
    text = re.sub(r'[🤖🎭🎤💾✅🔍💾📸📅📋🚀⚙️📖🔧❓⚡🗺️📄🙏📧]', '', text)
    
    # 如果没有特殊格式，直接添加
    if '`' not in text and '**' not in text and '[' not in text:
        run = paragraph.add_run(text)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        return
    
    # 处理格式：先处理代码，再处理粗体
    # 分割文本，按代码、粗体、普通文本处理
    parts = []
    last_end = 0
    
    # 先找到所有代码和粗体的位置
    matches = []
    for match in re.finditer(r'`([^`]+)`', text):
        matches.append(('code', match.start(), match.end(), match.group(1)))
    for match in re.finditer(r'\*\*([^*]+)\*\*', text):
        matches.append(('bold', match.start(), match.end(), match.group(1)))
    
    # 按位置排序
    matches.sort(key=lambda x: x[1])
    
    # 添加文本段
    for match_type, start, end, content in matches:
        if start > last_end:
            parts.append(('normal', text[last_end:start]))
        parts.append((match_type, content))
        last_end = end
    
    if last_end < len(text):
        parts.append(('normal', text[last_end:]))
    
    if not parts:
        parts = [('normal', text)]
    
    # 添加格式化的文本
    for part_type, part_text in parts:
        run = paragraph.add_run(part_text)
        
        if part_type == 'code':
            run.font.name = 'Consolas'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
            run.font.size = Pt(10)
            # 添加背景色
            shading = parse_xml(r'<w:shd {} w:fill="E8E8E8"/>'.format(nsdecls('w')))
            run._element.rPr.append(shading)
        elif part_type == 'bold':
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.bold = True
        else:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def convert_markdown_to_docx(md_file, docx_file):
    """将 Markdown 文件转换为 Word 文档"""
    print(f"正在读取: {md_file}")
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    doc = setup_document()
    
    # 添加标题
    title_para = doc.add_heading('玲 - 智能虚拟助手', 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加副标题
    subtitle_para = doc.add_paragraph('一个集成大语言模型、Live2D 虚拟形象、语音合成于一体的智能虚拟助手系统')
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.runs[0].font.italic = True
    
    in_code_block = False
    code_lines = []
    in_table = False
    table = None
    table_header_processed = False
    
    i = 0
    while i < len(lines):
        line = lines[i]
        line_type, content = parse_markdown_line(line)
        
        # 处理代码块
        if line_type == 'code_block':
            if in_code_block:
                # 结束代码块
                if code_lines:
                    code_para = doc.add_paragraph()
                    code_para.style = 'No Spacing'
                    code_run = code_para.add_run('\n'.join(code_lines))
                    code_run.font.name = 'Consolas'
                    code_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
                    code_run.font.size = Pt(10)
                    # 设置代码块背景
                    from docx.oxml.ns import nsdecls
                    shading_elm = parse_xml(r'<w:shd {} w:fill="F5F5F5"/>'.format(nsdecls('w')))
                    code_para._element.get_or_add_pPr().append(shading_elm)
                    code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line.rstrip())
            i += 1
            continue
        
        # 处理分隔线
        if line_type == 'hr':
            para = doc.add_paragraph('─' * 50)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # 处理空行
        if line_type == 'blank':
            doc.add_paragraph()
            i += 1
            continue
        
        # 处理标题
        if line_type.startswith('h'):
            level = int(line_type[1])
            heading = doc.add_heading(content, level)
            # 移除 emoji
            for run in heading.runs:
                run_text = run.text
                run_text = re.sub(r'[🤖🎭🎤💾✅🔍💾📸📅📋🚀⚙️📖🔧❓⚡🗺️📄🙏📧]', '', run_text)
                run.text = run_text
            i += 1
            continue
        
        # 处理表格（简单检测）
        if '|' in content and content.count('|') >= 2:
            cells = [c.strip() for c in content.split('|')[1:-1]]
            
            # 检查是否是分隔行（表头下方）
            if all(c.replace('-', '').replace(':', '').strip() == '' for c in cells):
                i += 1
                continue
            
            if not in_table:
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = 'Light Grid Accent 1'
                in_table = True
                table_header_processed = False
                row = table.rows[0]
                for j, cell_text in enumerate(cells):
                    if j < len(row.cells):
                        cell = row.cells[j]
                        cell.text = cell_text
                        # 设置表头格式
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                table_header_processed = True
            else:
                # 添加数据行
                if table:
                    row = table.add_row()
                    for j, cell_text in enumerate(cells):
                        if j < len(row.cells):
                            row.cells[j].text = cell_text
            i += 1
            continue
        
        # 重置表格状态
        if in_table and line_type != 'blank':
            in_table = False
            table = None
        
        # 处理列表
        if line_type == 'ul':
            para = doc.add_paragraph(style='List Bullet')
            add_formatted_text(para, content)
            i += 1
            continue
        
        if line_type == 'ol':
            para = doc.add_paragraph(style='List Number')
            add_formatted_text(para, content)
            i += 1
            continue
        
        # 处理普通段落
        if line_type == 'p':
            para = doc.add_paragraph()
            add_formatted_text(para, content)
            i += 1
            continue
        
        i += 1
    
    # 保存文档
    print(f"正在保存: {docx_file}")
    doc.save(docx_file)
    print(f"✅ 转换完成！")
    print(f"文件已保存到: {docx_file}")

def main():
    """主函数"""
    project_root = Path(__file__).parent
    md_file = project_root / "README.md"
    docx_file = project_root / "README.docx"
    
    if not md_file.exists():
        print(f"❌ 错误: 找不到文件 {md_file}")
        sys.exit(1)
    
    try:
        convert_markdown_to_docx(md_file, docx_file)
    except Exception as e:
        print(f"❌ 转换失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == '__main__':
    main()
